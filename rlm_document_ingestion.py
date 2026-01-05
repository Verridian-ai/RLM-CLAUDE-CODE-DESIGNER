#!/usr/bin/env python3
"""
Enhanced Document Ingestion Pipeline for RLM System
Supports multiple document formats with robust text extraction and structured output
"""

import os
import sys
import json
import mimetypes
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

# Set UTF-8 encoding for Windows console output
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

class DocumentIngestionPipeline:
    """
    Comprehensive document ingestion pipeline that handles multiple formats
    and provides RLM-compatible output
    """

    def __init__(self, output_dir: str = "ingested_documents"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

        # Document type handlers
        self.handlers = {
            'application/pdf': self._extract_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx,
            'text/html': self._extract_html,
            'text/plain': self._extract_text,
            'application/json': self._extract_json,
            'text/markdown': self._extract_markdown,
        }

    def detect_document_type(self, file_path: str) -> str:
        """Detect document type using multiple methods"""
        mime_type, _ = mimetypes.guess_type(file_path)

        if mime_type:
            return mime_type

        # Fallback: check file extension
        ext = Path(file_path).suffix.lower()
        ext_mapping = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.html': 'text/html',
            '.htm': 'text/html',
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.json': 'application/json',
            '.py': 'text/plain',
            '.js': 'text/plain',
            '.ts': 'text/plain',
            '.java': 'text/plain',
        }

        return ext_mapping.get(ext, 'application/octet-stream')

    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using multiple methods"""
        methods = [
            ('pypdf', self._pdf_pypdf),
            ('pdfplumber', self._pdf_pdfplumber),
            ('pymupdf', self._pdf_pymupdf),
        ]

        best_result = None
        best_score = 0

        for method_name, extract_func in methods:
            try:
                text, metadata = extract_func(file_path)
                if text:
                    score = len(text.strip())
                    if score > best_score:
                        best_score = score
                        best_result = (text, {**metadata, 'extraction_method': method_name})
            except Exception as e:
                print(f"  PDF method {method_name} failed: {e}")
                continue

        if best_result:
            return best_result
        else:
            raise ValueError("All PDF extraction methods failed")

    def _pdf_pypdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract using pypdf"""
        import pypdf
        text = ""
        reader = pypdf.PdfReader(file_path)

        for page_num, page in enumerate(reader.pages):
            text += f"\n--- PAGE {page_num + 1} ---\n"
            text += page.extract_text()

        metadata = {
            "page_count": len(reader.pages),
            "document_info": dict(reader.metadata) if reader.metadata else {},
        }

        return text, metadata

    def _pdf_pdfplumber(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract using pdfplumber"""
        import pdfplumber
        text = ""

        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text += f"\n--- PAGE {page_num + 1} ---\n"
                page_text = page.extract_text()
                if page_text:
                    text += page_text

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    text += f"\n[TABLES ON PAGE {page_num + 1}]\n"
                    for table_num, table in enumerate(tables):
                        text += f"\nTable {table_num + 1}:\n"
                        for row in table:
                            text += " | ".join([cell or "" for cell in row]) + "\n"

            metadata = {
                "page_count": len(pdf.pages),
                "document_info": dict(pdf.metadata) if pdf.metadata else {},
            }

        return text, metadata

    def _pdf_pymupdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract using PyMuPDF"""
        import fitz
        text = ""
        doc = fitz.open(file_path)

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += f"\n--- PAGE {page_num + 1} ---\n"
            text += page.get_text()

        metadata = {
            "page_count": len(doc),
            "document_info": dict(doc.metadata),
        }

        doc.close()
        return text, metadata

    def _extract_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX files"""
        try:
            import docx
            doc = docx.Document(file_path)

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                text += "\n[TABLE]\n"
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text += " | ".join(cells) + "\n"

            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            return text, metadata

        except ImportError:
            raise ValueError("python-docx library not available for DOCX extraction")

    def _extract_html(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from HTML files"""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            text = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            text = '\n'.join(line for line in lines if line)

            metadata = {
                "title": soup.title.string if soup.title else "",
                "links_count": len(soup.find_all('a')),
                "images_count": len(soup.find_all('img')),
            }

            return text, metadata

        except ImportError:
            raise ValueError("BeautifulSoup library not available for HTML extraction")

    def _extract_text(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()

                metadata = {
                    "encoding_used": encoding,
                    "character_count": len(text),
                    "line_count": text.count('\n'),
                }

                return text, metadata

            except UnicodeDecodeError:
                continue

        raise ValueError("Could not decode text file with any common encoding")

    def _extract_json(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from JSON files"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Convert JSON to readable text
        text = json.dumps(data, indent=2, ensure_ascii=False)

        metadata = {
            "json_structure": self._analyze_json_structure(data),
            "total_keys": self._count_json_keys(data),
        }

        return text, metadata

    def _extract_markdown(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from Markdown files"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()

        # Basic markdown analysis
        headers = [line for line in text.split('\n') if line.strip().startswith('#')]

        metadata = {
            "header_count": len(headers),
            "headers": headers[:10],  # First 10 headers
            "estimated_words": len(text.split()),
        }

        return text, metadata

    def _analyze_json_structure(self, data: Any, depth: int = 0) -> Dict[str, Any]:
        """Analyze JSON structure recursively"""
        if depth > 5:  # Prevent infinite recursion
            return {"type": type(data).__name__, "truncated": True}

        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys())[:10],  # First 10 keys
                "key_count": len(data),
                "sample_values": {k: self._analyze_json_structure(v, depth + 1)
                                 for k, v in list(data.items())[:3]}
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "sample_items": [self._analyze_json_structure(item, depth + 1)
                               for item in data[:3]]
            }
        else:
            return {"type": type(data).__name__, "value": str(data)[:100]}

    def _count_json_keys(self, data: Any) -> int:
        """Count total keys in nested JSON"""
        if isinstance(data, dict):
            return len(data) + sum(self._count_json_keys(v) for v in data.values())
        elif isinstance(data, list):
            return sum(self._count_json_keys(item) for item in data)
        else:
            return 0

    def create_rlm_compatible_output(self, text: str, metadata: Dict[str, Any],
                                   file_path: str) -> Dict[str, Any]:
        """Create RLM-compatible structured output"""
        file_name = Path(file_path).stem

        # Estimate token count (rough approximation)
        token_count = len(text.split()) * 1.3  # Average tokens per word

        # Create chunk boundaries if text is large
        chunk_size = 10000  # Characters per chunk
        chunks = []
        if len(text) > chunk_size:
            for i in range(0, len(text), chunk_size):
                chunk_text = text[i:i + chunk_size]
                chunks.append({
                    "chunk_id": f"{file_name}_chunk_{len(chunks) + 1}",
                    "start_char": i,
                    "end_char": min(i + chunk_size, len(text)),
                    "preview": chunk_text[:200] + "..." if len(chunk_text) > 200 else chunk_text
                })

        rlm_output = {
            "document_id": file_name,
            "source_file": str(Path(file_path).absolute()),
            "ingestion_timestamp": datetime.now().isoformat(),
            "document_type": metadata.get('extraction_method', 'unknown'),
            "content": text,
            "metadata": metadata,
            "statistics": {
                "character_count": len(text),
                "estimated_token_count": int(token_count),
                "chunk_count": len(chunks),
                "requires_chunking": len(text) > chunk_size
            },
            "chunks": chunks,
            "rlm_processing_hints": {
                "recommended_chunk_size": chunk_size,
                "processing_method": "recursive_delegation" if len(text) > 50000 else "direct_processing",
                "complexity_score": min(10, max(1, len(text) // 10000))
            }
        }

        return rlm_output

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a single document through the ingestion pipeline"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")

        print(f"Processing document: {file_path}")

        # Detect document type
        doc_type = self.detect_document_type(file_path)
        print(f"  Detected type: {doc_type}")

        # Extract content
        if doc_type in self.handlers:
            try:
                text, metadata = self.handlers[doc_type](file_path)
                print(f"  Extracted {len(text)} characters")
            except Exception as e:
                print(f"  Extraction failed: {e}")
                raise
        else:
            # Fallback to text extraction
            try:
                text, metadata = self._extract_text(file_path)
                print(f"  Used fallback text extraction: {len(text)} characters")
            except Exception as e:
                raise ValueError(f"Unsupported document type: {doc_type}, fallback failed: {e}")

        # Create RLM-compatible output
        rlm_output = self.create_rlm_compatible_output(text, metadata, file_path)

        # Save structured output
        output_file = self.output_dir / f"{Path(file_path).stem}_ingested.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(rlm_output, f, indent=2, ensure_ascii=False)

        # Save raw text for RLM processing
        text_file = self.output_dir / f"{Path(file_path).stem}_text.txt"
        with open(text_file, 'w', encoding='utf-8') as f:
            f.write(text)

        print(f"  Saved to: {output_file}")
        print(f"  Text file: {text_file}")

        return rlm_output

def main():
    if len(sys.argv) != 2:
        print("Usage: python rlm_document_ingestion.py <document_path>")
        sys.exit(1)

    document_path = sys.argv[1]

    try:
        pipeline = DocumentIngestionPipeline()
        result = pipeline.process_document(document_path)

        print("\nIngestion Summary:")
        print("=" * 50)
        print(f"Document ID: {result['document_id']}")
        print(f"Type: {result['document_type']}")
        print(f"Characters: {result['statistics']['character_count']:,}")
        print(f"Estimated Tokens: {result['statistics']['estimated_token_count']:,}")
        print(f"Chunks: {result['statistics']['chunk_count']}")
        print(f"RLM Processing: {result['rlm_processing_hints']['processing_method']}")
        print(f"Complexity Score: {result['rlm_processing_hints']['complexity_score']}/10")

        return True

    except Exception as e:
        print(f"Error: {e}")
        return False

if __name__ == "__main__":
    main()